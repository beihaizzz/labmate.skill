#!/usr/bin/env python3
"""Auto-prepare blank template — detect label cells and inject {{placeholder}} into adjacent fillable cells.

GREEN phase: horizontal detection only. Scans tables for label cells (detected via
is_label_cell_v2), then injects {{role}} placeholders into adjacent fillable cells
in the same row. Original input file is never modified — always operates on a copy.

Vertical adjacency (T5) and full hint text regex replacement (T6) are deferred.
"""

import argparse
import json
import re
import sys
import shutil
import tempfile
import hashlib
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

# Import from same-directory modules
sys.path.insert(0, str(Path(__file__).parent))
import fill_utils
from role_aliases import ROLE_ALIASES, is_label_cell_v2, is_hint_text, _normalize_role


# ── PrepareResult ──────────────────────────────────────────────────────────────────

class PrepareResult(dict):
    """Dict-like result whose __str__ returns the output file path.

    This allows tests to use `Document(str(result))` to open the output file
    while also checking `isinstance(result, dict)` for dict-based assertions.
    """
    def __str__(self):
        return self.get("output", "")


# ── Fillable cell detection ───────────────────────────────────────────────────────

def _is_fillable(cell_text: str) -> bool:
    """Check if a cell is fillable (empty or contains hint/placeholder text)."""
    t = cell_text.strip()
    if not t:
        return True
    if is_hint_text(t):
        return True
    return False


# ── Placeholder already present? ──────────────────────────────────────────────────

_PLACEHOLDER_RE = re.compile(r'\{\{[^}]+\}\}')


def _has_placeholder(text: str) -> bool:
    """Check if text already contains a {{placeholder}} marker."""
    return bool(_PLACEHOLDER_RE.search(text))


# ── Format capture ────────────────────────────────────────────────────────────────

def _capture_cell_format(cell) -> dict:
    """Capture formatting from the first run of the first paragraph in a cell.

    Returns a dict with font_name, font_size_pt, bold, east_asia, alignment.
    All values are None if no formatting is found.
    """
    fmt = {
        "font_name": None,
        "font_size_pt": None,
        "bold": False,
        "east_asia": None,
        "alignment": None,
    }

    first_para = cell.paragraphs[0]
    if first_para.alignment is not None:
        fmt["alignment"] = str(first_para.alignment)

    for run in first_para.runs:
        if run.font.name:
            fmt["font_name"] = run.font.name
        if run.font.size:
            fmt["font_size_pt"] = round(run.font.size / 12700, 1)
        if run.font.bold is not None:
            fmt["bold"] = run.font.bold

        # Read eastAsia from XML directly
        rPr = run._element.find(qn('w:rPr'))
        if rPr is not None:
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                ea = rFonts.get(qn('w:eastAsia'))
                if ea:
                    fmt["east_asia"] = ea
        break  # Only capture first run

    return fmt


# ── Placeholder injection ─────────────────────────────────────────────────────────

def _inject_placeholder(cell, role: str, fmt: dict):
    """Clear a cell and inject a {{role}} placeholder with captured formatting.

    Args:
        cell: python-docx Cell object.
        role: Normalized role name (e.g., "课程名称").
        fmt: Formatting dict from _capture_cell_format().
    """
    # Clear existing content
    for p in cell.paragraphs:
        p.clear()

    first_para = cell.paragraphs[0]
    placeholder = "{{" + role + "}}"
    new_run = first_para.add_run(placeholder)

    # Re-apply captured formatting
    fill_utils.set_run_font(
        new_run,
        font_name=fmt.get("font_name"),
        font_size_pt=fmt.get("font_size_pt") or fill_utils.SIZE_BODY,
        bold=fmt.get("bold"),
        east_asia=fmt.get("east_asia"),
    )

    if fmt.get("alignment"):
        fill_utils.set_paragraph_alignment(first_para, fmt["alignment"])


# ── Main prepare function ─────────────────────────────────────────────────────────

def prepare_template(input_path: str, output_path: str = None) -> PrepareResult:
    """Auto-prepare a blank template by injecting placeholders into fillable cells.

    Args:
        input_path: Path to the blank template DOCX file.
        output_path: Path to write the prepared output. If None, a temp file is used.

    Returns:
        PrepareResult (dict subclass) with keys:
            success: bool
            source: input file path
            output: output file path
            roles_injected: list of dicts with role, table, row, col, cell_text_was
            warnings: list of warning strings
    """
    source = Path(input_path).resolve()

    # If no output path given, create a temp file
    if output_path is None:
        tmp_fd, tmp_path = tempfile.mkstemp(suffix='.docx', prefix='auto_prepared_')
        import os
        os.close(tmp_fd)
        output_path = tmp_path

    output = Path(output_path)

    # Always operate on a copy — original stays untouched
    shutil.copy(source, output)

    roles_injected = []
    warnings = []

    try:
        doc = Document(str(output))

        for t_idx, table in enumerate(doc.tables):
            # ── Horizontal pass ──────────────────────────────────────────────
            for r_idx, row in enumerate(table.rows):
                # Phase 1: Find label cells in this row
                label_positions = []  # (col_idx, role)

                for c_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if not text:
                        continue

                    # NOTE: is_label_cell_v2() internally checks is_hint_text() and
                    # returns False for hint patterns like "_____", "（请填写）", "待定".
                    # This is critical: "待定" is 2-char CJK and would otherwise pass
                    # the label heuristics (length ≤ 12, contains CJK, no bullet prefix),
                    # but is_hint_text() catches it first.  Hint cells are then detected
                    # as fillable by _is_fillable() and receive {{role}} placeholders.
                    if is_label_cell_v2(text):
                        role = _normalize_role(text)
                        label_positions.append((c_idx, role))

                if not label_positions:
                    continue  # No labels in this row, nothing to match

                # Phase 2: For each fillable cell, assign nearest label
                for c_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()

                    # Skip cells that already have placeholders
                    if _has_placeholder(cell_text):
                        continue

                    # Only target fillable cells
                    if not _is_fillable(cell_text):
                        continue

                    # Skip if this cell IS a label itself (shouldn't happen but guard)
                    is_label = is_label_cell_v2(cell_text)
                    if is_label:
                        continue

                    # Find nearest label by column distance
                    nearest_label = min(
                        label_positions,
                        key=lambda lp: abs(lp[0] - c_idx)
                    )
                    nearest_role = nearest_label[1]

                    # Capture formatting before clearing
                    fmt = _capture_cell_format(cell)

                    # Inject placeholder
                    _inject_placeholder(cell, nearest_role, fmt)

                    roles_injected.append({
                        "role": nearest_role,
                        "direction": "horizontal",
                        "table": t_idx,
                        "row": r_idx,
                        "col": c_idx,
                        "cell_text_was": cell_text or "(empty)",
                    })

            # ── Vertical pass ────────────────────────────────────────────────
            max_row = len(table.rows) - 1
            for r_idx, row in enumerate(table.rows):
                # Skip the last row — no row below to inject into
                if r_idx >= max_row:
                    continue

                for c_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if not text:
                        continue

                    # is_label_cell_v2() internally rejects hint text via is_hint_text(),
                    # so "待定" / "_____" / "（请填写）" are NOT classified as labels.
                    if not is_label_cell_v2(text):
                        continue

                    # This is a label cell — check the cell directly below
                    role = _normalize_role(text)
                    below_cell = table.rows[r_idx + 1].cells[c_idx]
                    below_text = below_cell.text.strip()

                    # Must be fillable (empty or hint text)
                    if not _is_fillable(below_text):
                        continue

                    # Must not be a label itself
                    if is_label_cell_v2(below_text):
                        continue

                    # Must not already have a placeholder (horizontal may have filled it)
                    if _has_placeholder(below_text):
                        continue

                    # Inject placeholder vertically
                    fmt = _capture_cell_format(below_cell)
                    _inject_placeholder(below_cell, role, fmt)

                    roles_injected.append({
                        "role": role,
                        "direction": "vertical",
                        "table": t_idx,
                        "row": r_idx + 1,
                        "col": c_idx,
                        "cell_text_was": below_text or "(empty)",
                    })

        # Special handling for merged cells that span across and contain headers
        # (e.g., "实验报告封面" spanning 4 cols) — these are >12 chars so
        # is_label_cell_v2 already rejects them. No special handling needed.

        doc.save(str(output))

        result = PrepareResult({
            "success": True,
            "source": str(source),
            "output": str(output),
            "roles_injected": roles_injected,
            "warnings": warnings,
        })
        return result

    except Exception as e:
        return PrepareResult({
            "success": False,
            "source": str(source),
            "output": str(output),
            "roles_injected": roles_injected,
            "warnings": warnings + [f"Error during preparation: {e}"],
        })


# ── CLI ────────────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description='Auto-prepare blank template — inject {{placeholders}} into fillable cells'
    )
    parser.add_argument('--input', '-i', required=True,
                        help='Path to blank template .docx file')
    parser.add_argument('--output', '-o', required=True,
                        help='Path to write the prepared output .docx')
    parser.add_argument('--format', '-f', choices=['json', 'human'], default='json',
                        help='Output format (default: json)')
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: Input file not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    result = prepare_template(str(input_path), str(Path(args.output)))

    if args.format == 'human':
        if result["success"]:
            print(f"✅ Template prepared successfully")
            print(f"   Source: {result['source']}")
            print(f"   Output: {result['output']}")
            print(f"   Roles injected: {len(result['roles_injected'])}")
            for entry in result["roles_injected"]:
                print(f"     - Table {entry['table']} R{entry['row']}C{entry['col']}: "
                      f"{{{{{entry['role']}}}}} (was: {entry['cell_text_was']})")
            for w in result["warnings"]:
                print(f"   ⚠️  {w}")
        else:
            print(f"❌ Preparation failed")
            for w in result["warnings"]:
                print(f"   Error: {w}")
    else:
        # JSON output
        print(json.dumps(dict(result), indent=2, ensure_ascii=False))

    sys.exit(0 if result.get("success") else 1)


if __name__ == '__main__':
    main()
