#!/usr/bin/env python3
"""Word document parsing (.doc and .docx) with placeholder detection."""

import argparse
import json
import re
import shutil
import subprocess
import sys
from pathlib import Path

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def _find_libreoffice() -> str | None:
    """Find LibreOffice executable."""
    for name in ['soffice', 'libreoffice', 'soffice.exe', 'libreoffice.exe']:
        found = shutil.which(name)
        if found:
            return found
    # Check common Windows install paths
    for base in [r'C:\Program Files\LibreOffice\program',
                 r'C:\Program Files (x86)\LibreOffice\program']:
        for exe in ['soffice.exe', 'swriter.exe']:
            cand = Path(base) / exe
            if cand.exists():
                return str(cand)
    return None


def _convert_to_docx(filepath: Path) -> Path | None:
    """Convert .doc to .docx. Returns the .docx path or None."""
    lo = _find_libreoffice()
    if not lo:
        return None
    try:
        result = subprocess.run(
            [lo, '--headless', '--convert-to', 'docx', '--outdir',
             str(filepath.parent), str(filepath)],
            capture_output=True, text=True, timeout=60
        )
        if result.returncode == 0:
            docx = filepath.with_suffix('.docx')
            if docx.exists():
                return docx
    except Exception:
        pass
    return None


def parse_docx(filepath: Path):
    """Parse DOCX and extract text, tables, placeholders."""
    if not HAS_DOCX:
        return {"error": "Missing dependency: python-docx"}

    suffix = filepath.suffix.lower()
    
    # .doc → auto-convert
    if suffix == '.doc':
        docx_path = _convert_to_docx(filepath)
        if docx_path is None:
            return {
                "error": (
                    "Cannot read .doc file. Please:\n"
                    "1. Install LibreOffice (https://www.libreoffice.org/)\n"
                    "2. Or manually save as .docx in Word / WPS\n"
                    "3. Or provide a .docx version of this file"
                ),
                "needs_conversion": True,
                "filename": filepath.name
            }
        filepath = docx_path

    result = {
        "filename": filepath.name,
        "paragraphs": [],
        "tables": [],
        "placeholders": [],
        "structure": {
            "paragraph_count": 0,
            "table_count": 0,
            "heading_count": 0
        }
    }

    try:
        doc = Document(filepath)

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                result["paragraphs"].append({
                    "text": text,
                    "style": para.style.name if para.style else None
                })
                if para.style and "heading" in para.style.name.lower():
                    result["structure"]["heading_count"] += 1

        full_text = "\n".join([p["text"] for p in result["paragraphs"]])
        placeholders = re.findall(r'\{\{([^}]+)\}\}', full_text)
        result["placeholders"] = list(set([f"{{{{{p}}}}}" for p in placeholders]))

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                    cell_placeholders = re.findall(r'\{\{([^}]+)\}\}', cell_text)
                    for p in cell_placeholders:
                        ph = f"{{{{{p}}}}}"
                        if ph not in result["placeholders"]:
                            result["placeholders"].append(ph)
                table_data.append(row_data)
            result["tables"].append(table_data)

        result["structure"]["paragraph_count"] = len(result["paragraphs"])
        result["structure"]["table_count"] = len(result["tables"])

    except Exception as e:
        result["error"] = str(e)

    return result


def main():
    parser = argparse.ArgumentParser(description='Parse Word document (.doc / .docx)')
    parser.add_argument('--input', '-i', required=True, help='Input .doc or .docx file')
    args = parser.parse_args()

    filepath = Path(args.input)
    if not filepath.exists():
        print(f"Error: File not found: {filepath}", file=sys.stderr)
        sys.exit(1)

    result = parse_docx(filepath)
    print(json.dumps(result, indent=2, ensure_ascii=False))

    sys.exit(0 if "error" not in result else 1)


if __name__ == '__main__':
    main()
