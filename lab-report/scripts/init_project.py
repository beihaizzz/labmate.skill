#!/usr/bin/env python3
"""Project initialization orchestration."""

import argparse
import json
import subprocess
import sys
from pathlib import Path

DEFAULT_WORKING_DIR = ".labmate"

# Import other scripts
sys.path.insert(0, str(Path(__file__).parent))

def run_check_deps():
    """Run dependency check."""
    try:
        result = subprocess.run(
            [sys.executable, str(Path(__file__).parent / "check_deps.py")],
            capture_output=True,
            text=True
        )
        return result.returncode == 0, result.stdout
    except Exception as e:
        return False, str(e)

def discover_files(directory: Path) -> dict:
    """Discover course materials in directory."""
    files = {
        "guides": [],
        "templates": [],
        "references": []
    }
    
    for file in directory.iterdir():
        if not file.is_file():
            continue
        
        suffix = file.suffix.lower()
        
        if suffix == '.pdf':
            files["guides"].append(file.name)
        elif suffix in ['.docx', '.doc']:
            files["templates"].append(file.name)
        elif suffix == '.pptx':
            files["guides"].append(file.name)
        elif suffix in ['.md', '.txt', '.cs', '.py', '.cpp', '.h']:
            files["references"].append(file.name)
    
    # Scan subdirectories for reference code folders
    for sub in directory.iterdir():
        if not sub.is_dir():
            continue
        name_lower = sub.name.lower()
        if any(kw in name_lower for kw in ['供参考', '参考', 'reference', 'script', 'scripts', '资源']):
            count = len(list(sub.rglob('*')))
            if count > 0:
                files.setdefault("reference_dirs", [])
                files["reference_dirs"].append({"name": sub.name, "file_count": count})
    
    return files


TEMPLATE_FINGERPRINTS = [
    "课程名称", "学生姓名", "学号", "专业年级", "任课教师",
    "实验名称", "实验类型", "实验学时", "实验日期", "实验地点",
    "实验目的", "实验原理", "实验内容", "实验要求",
]


def _extract_text_from_doc(filepath: Path) -> str:
    """从 .doc (OLE2) 提取纯文本。（无 LibreOffice 降级方案）"""
    try:
        import olefile
        ole = olefile.OleFileIO(str(filepath))
        # 尝试读取 WordDocument 流
        if ole.exists('WordDocument'):
            stream = ole.openstream('WordDocument')
            raw = stream.read()
            # 过滤可读字符
            text = ''.join(chr(b) for b in raw if 0x20 <= b < 0x7f or 0x80 <= b <= 0xff)
            # 清理乱码序列
            text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)
            return text
        ole.close()
    except ImportError:
        pass
    except Exception:
        pass
    return ""


def _detect_embedded_template(filepath: Path) -> dict | None:
    """检测 .doc 文件中是否嵌有报告模板。"""
    suffix = filepath.suffix.lower()
    
    if suffix == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            text = " ".join([p.text for p in doc.paragraphs])
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += " " + cell.text
        except Exception:
            return None
    elif suffix == '.doc':
        text = _extract_text_from_doc(filepath)
    else:
        return None
    
    if not text:
        return None
    
    found = [fp for fp in TEMPLATE_FINGERPRINTS if fp in text]
    if len(found) >= 4:
        return {"source": filepath.name, "cells_detected": found,
                "type": "embedded_in_guide", "status": "detected"}
    return None


def _save_config(directory: Path, data: dict, working_dir: str = DEFAULT_WORKING_DIR):
    """保存配置到 working_dir/config.json"""
    config_dir = directory / working_dir
    config_dir.mkdir(exist_ok=True)
    config_file = config_dir / "config.json"
    existing = {}
    if config_file.exists():
        try:
            existing = json.loads(config_file.read_text(encoding='utf-8'))
        except (json.JSONDecodeError, IOError):
            pass
    existing.update(data)
    config_file.write_text(json.dumps(existing, indent=2, ensure_ascii=False), encoding='utf-8')


def _create_or_update_project_md(directory: Path, experiment_name: str = None,
                                  course_info: dict = None,
                                  student_info: dict = None) -> Path:
    """在课程根目录创建或更新 project.md。

    project.md 是课程级别的上下文文件，记录课程信息和实验进度。
    不存敏感信息（学生姓名等存于学生信息.md），仅存元信息。
    """
    project_md = directory / "project.md"

    # 解析已有内容（只读 key: value 行）
    existing = {}
    if project_md.exists():
        for line in project_md.read_text(encoding='utf-8').split('\n'):
            if ':' in line and not line.startswith(('#', '- ', '>')):
                k, v = line.split(':', 1)
                existing[k.strip()] = v.strip()
    if course_info:
        existing.update(course_info)

    lines = ["# 课程信息\n"]
    for key in ["课程名称", "课程代码", "任课教师"]:
        lines.append(f"{key}: {existing.get(key, f'{{{{key}}}}')}")

    # 实验进度（保留已有 checkbox，追加新实验）
    lines.append("\n# 实验进度\n")
    old_boxes = set()
    if project_md.exists():
        for line in project_md.read_text(encoding='utf-8').split('\n'):
            if line.startswith('- [') and '实验' in line:
                old_boxes.add(line.strip())
                lines.append(line)
    if experiment_name:
        entry = f"- [ ] {experiment_name}"
        if entry not in old_boxes:
            lines.append(entry)
    if not any('- [' in l for l in lines if l.startswith('- [')):
        lines.append("- [ ] ...")

    lines.append("\n# 通用配置\n")
    lines.append("默认风格: normal")
    lines.append("Git: 未启用")
    lines.append("\n---\n> 由 lab-report skill 自动维护。AI 每次会话启动时先读此文件了解项目状态。")

    project_md.write_text('\n'.join(lines), encoding='utf-8')
    return project_md

def init_project(directory: Path, use_git: bool = False, experiment_name: str = None):
    """Initialize project."""
    result = {
        "success": True,
        "directory": str(directory),
        "discovered_files": {},
        "student_info": None,
        "progress_initialized": False,
        "git_initialized": False,
        "errors": []
    }
    
    # 1. Check dependencies
    deps_ok, deps_output = run_check_deps()
    if not deps_ok:
        result["errors"].append(f"Dependency check failed: {deps_output}")
        result["success"] = False
        return result
    
    # 2. Discover files
    result["discovered_files"] = discover_files(directory)
    
    # 2.5 检测 .doc 嵌入模板
    result["embedded_templates"] = []
    for file in directory.iterdir():
        if not file.is_file():
            continue
        detected = _detect_embedded_template(file)
        if detected:
            result["embedded_templates"].append(detected)
            result["discovered_files"].setdefault("templates_embedded", [])
            result["discovered_files"]["templates_embedded"].append(file.name)
    
    if result["embedded_templates"]:
        _save_config(directory, {"embedded_templates": result["embedded_templates"]}, working_dir)

    if not any(result["discovered_files"].values()):
        result["errors"].append("No course materials found in directory")
        result["success"] = False
        return result

    # 4. Check student info
    try:
        from student_info import find_student_info
        info_path, info_data = find_student_info(directory)
        result["student_info"] = {
            "found": info_path is not None,
            "path": str(info_path) if info_path else None,
            "data": info_data
        }
    except Exception as e:
        result["errors"].append(f"Student info check failed: {e}")

    # 5. Determine working directory: use existing .lab-report/ if present, else .labmate/
    if (directory / ".lab-report").exists():
        working_dir = ".lab-report"
    else:
        working_dir = DEFAULT_WORKING_DIR

    lab_dir = directory / working_dir
    lab_dir.mkdir(parents=True, exist_ok=True)

    # Save working_dir to config
    _save_config(directory, {"working_dir": working_dir}, working_dir)
    
    # 6. Initialize progress if experiment name provided
    if experiment_name:
        try:
            # Count steps from PDF if available
            total_steps = 5  # Default
            
            from progress_manager import init_progress
            import os
            os.chdir(directory)
            init_progress(experiment_name, total_steps)
            result["progress_initialized"] = True
        except Exception as e:
            result["errors"].append(f"Progress init failed: {e}")
    
    # 7. Initialize git if requested
    if use_git:
        try:
            git_dir = directory / ".git"
            if not git_dir.exists():
                subprocess.run(
                    ["git", "init"],
                    cwd=directory,
                    capture_output=True,
                    check=True
                )
                subprocess.run(
                    ["git", "add", "."],
                    cwd=directory,
                    capture_output=True
                )
                subprocess.run(
                    ["git", "commit", "-m", "Initial commit"],
                    cwd=directory,
                    capture_output=True
                )
                result["git_initialized"] = True
        except Exception as e:
            result["errors"].append(f"Git init failed: {e}")
    
    return result

def main():
    parser = argparse.ArgumentParser(description='Initialize lab-report project')
    parser.add_argument('--dir', type=Path, default=Path.cwd(), help='Target directory')
    parser.add_argument('--git', action='store_true', help='Initialize git repository')
    parser.add_argument('--name', help='Experiment name')
    args = parser.parse_args()
    
    if not args.dir.exists():
        print(f"Error: Directory not found: {args.dir}", file=sys.stderr)
        sys.exit(1)
    
    result = init_project(args.dir, args.git, args.name)
    print(json.dumps(result, indent=2, ensure_ascii=False))
    
    sys.exit(0 if result["success"] else 1)

if __name__ == '__main__':
    main()
