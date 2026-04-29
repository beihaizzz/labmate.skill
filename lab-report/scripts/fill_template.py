#!/usr/bin/env python3
"""Template filling — uses inspect data for exact formatting (no guessing)."""

import argparse
import hashlib
import json
import re
import shutil
import subprocess
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docxtpl import DocxTemplate
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# Import formatting utilities (outside try — fill_utils has its own import guards)
sys.path.insert(0, str(Path(__file__).parent))
import fill_utils

BANNED_WORDS = ['首先', '其次', '最后', '总而言之', '值得注意的是', '综上所述', '不可否认']


def _find_libreoffice() -> str | None:
    for name in ['soffice', 'libreoffice', 'soffice.exe', 'libreoffice.exe']:
        found = shutil.which(name)
        if found:
            return found
    for base in [r'C:\Program Files\LibreOffice\program',
                 r'C:\Program Files (x86)\LibreOffice\program']:
        for exe in ['soffice.exe', 'swriter.exe']:
            cand = Path(base) / exe
            if cand.exists():
                return str(cand)
    return None


def _convert_to_docx(filepath: Path) -> Path | None:
    lo = _find_libreoffice()
    if not lo:
        return None
    try:
        result = subprocess.run(
            [lo, '--headless', '--convert-to', 'docx', '--outdir',
             str(filepath.parent), str(filepath)],
            capture_output=True, text=True, timeout=60)
        if result.returncode == 0:
            docx = filepath.with_suffix('.docx')
            if docx.exists():
                return docx
    except Exception:
        pass
    return None


def _verify_no_missing_placeholders(doc: Document) -> list:
    full_text = " ".join([p.text for p in doc.paragraphs])
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += " " + cell.text
    return re.findall(r'\{\{([^}]+)\}\}', full_text)


def _compare_with_inspect(output_doc: Document, inspect_data: dict) -> list:
    """Post-fill diff: detect overwritten label cells."""
    warnings = []
    out_tables = output_doc.tables
    for t_idx, table in enumerate(inspect_data.get("tables", [])):
        if t_idx >= len(out_tables):
            break
        for cell_info in table.get("cells", []):
            if not cell_info.get("is_label"):
                continue
            out_cell = out_tables[t_idx].rows[cell_info["row"]].cells[cell_info["column"]]
            out_text = out_cell.text.strip()
            in_text = cell_info.get("text", "").strip()
            if out_text and out_text != in_text and in_text:
                warnings.append(
                    f"LABEL R{cell_info['row']}C{cell_info['column']}: "
                    f"\"{in_text}\" overridden as \"{out_text[:50]}\""
                )
    return warnings


def _build_cell_index(inspect_data: dict):
    """From inspect JSON, build: cell_ref[(r,c)]=cell_info, label_cells set, placeholder_map."""
    cell_ref = {}
    label_cells = set()
    placeholder_map = {}
    if not inspect_data:
        return cell_ref, label_cells, placeholder_map
    for table in inspect_data.get("tables", []):
        for ci in table.get("cells", []):
            key = (ci["row"], ci["column"])
            cell_ref[key] = ci
            if ci.get("is_label"):
                label_cells.add(key)
            for p in ci.get("paragraphs", []):
                for rn in p.get("runs", []):
                    for ph in re.findall(r'\{\{[^}]+\}\}', rn.get("text_preview", "")):
                        if ph not in placeholder_map:
                            placeholder_map[ph] = key
    return cell_ref, label_cells, placeholder_map


def _get_fmt_from_ref(cell_ref_key, row, col, inspect_data):
    """Extract font_name, size_pt, bold, east_asia, alignment from the first non-empty run."""
    ci = None
    for table in inspect_data.get("tables", []):
        for c in table.get("cells", []):
            if c["row"] == row and c["column"] == col:
                ci = c
                break
        if ci:
            break
    if not ci:
        return None, None, None, None, None

    # Get alignment from first paragraph
    alignment = None
    for p in ci.get("paragraphs", []):
        if p.get("alignment"):
            alignment = p["alignment"]
            break

    # Get font info from first non-empty run
    for p in ci.get("paragraphs", []):
        for rn in p.get("runs", []):
            if rn.get("text_preview", "").strip():
                return (rn.get("font_name"),
                        rn.get("font_size_pt"),
                        rn.get("bold"),
                        rn.get("east_asia"),
                        alignment)

    return None, None, None, None, None


def fill_with_inspect(template_path: Path, data_path: Path, output_path: Path,
                      inspect_data: dict = None):
    """Fill template using inspect data for exact formatting."""
    if not HAS_DOCX:
        return {"error": "Missing dependencies: python-docx, docxtpl"}

    result = {
        "success": False, "template": str(template_path), "output": str(output_path),
        "placeholders_filled": [], "placeholders_missing": [], "warnings": [],
    }

    cell_ref, label_cells, placeholder_map = _build_cell_index(inspect_data)

    suffix = template_path.suffix.lower()
    if suffix == '.doc':
        conv = _convert_to_docx(template_path)
        if not conv:
            result["error"] = "Cannot process .doc. Install LibreOffice or save as .docx."
            return result
        template_path = conv

    try:
        with open(data_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except Exception as e:
        result["error"] = f"Cannot read data: {e}"
        return result

    try:
        shutil.copy(template_path, output_path)
        doc = DocxTemplate(output_path)
        doc.render(data)
        doc.save(output_path)
    except Exception as e:
        result["error"] = f"docxtpl render failed: {e}"
        return result

    try:
        final_doc = Document(output_path)

        for t_idx, table in enumerate(final_doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    key = (r_idx, c_idx)

                    # PRESERVE label cells — do NOT modify (fix 2)
                    if inspect_data and key in label_cells:
                        orig = cell_ref[key].get("text", "")
                        cur = cell.text.strip()
                        if cur and cur != orig:
                            # Label was overwritten — flag warning
                            # Attempt to restore original text
                            cell.paragraphs[0].clear()
                            run = cell.paragraphs[0].add_run(orig)
                        continue

                    # Apply template-correct formatting from inspect data (fix 1, 5)
                    ref = cell_ref.get(key)
                    if ref and inspect_data:
                        fname, fsize, fbold, fea, falign = _get_fmt_from_ref(
                            ref, r_idx, c_idx, inspect_data)
                        for para in cell.paragraphs:
                            # Apply paragraph alignment (fix 2.1)
                            fill_utils.set_paragraph_alignment(para, falign)
                            # Apply first-line indent for body paragraphs (fix 2.2)
                            fill_utils.apply_first_line_indent(para)
                            for run in para.runs:
                                fill_utils.set_run_font(
                                    run, font_name=fname, font_size_pt=fsize,
                                    bold=fbold, east_asia=fea)

        final_doc.save(output_path)

        # Post-fill diff check
        if inspect_data:
            check = Document(output_path)
            result["warnings"] = _compare_with_inspect(check, inspect_data)

        missing = _verify_no_missing_placeholders(Document(output_path))
        result["placeholders_missing"] = missing
        result["success"] = True

    except Exception as e:
        result["error"] = f"Post-processing failed: {e}"

    return result


def fill_cells_direct(template_path: Path, cells_data: list, output_path: Path,
                       inspect_data: dict = None):
    """Direct table fill mode — for templates without {{placeholders}}.
    
    cells_data: [
      {"table": 0, "row": 3, "col": 1, "text": "张啸瑞",
       "font_name": "宋体", "font_size": 14, "bold": false, "align": "CENTER", "east_asia": null}, ...
    ]
    """
    if not HAS_DOCX:
        return {"error": "Missing dependencies: python-docx"}

    result = {"success": False, "template": str(template_path), "output": str(output_path),
              "cells_filled": 0, "warnings": []}

    suffix = template_path.suffix.lower()
    if suffix == '.doc':
        conv = _convert_to_docx(template_path)
        if not conv:
            result["error"] = "Cannot process .doc. Install LibreOffice or save as .docx."
            return result
        template_path = conv

    try:
        shutil.copy(template_path, output_path)
        doc = Document(output_path)

        for entry in cells_data:
            t = entry.get("table", 0)
            r = entry.get("row", 0)
            c = entry.get("col", 0)
            if t >= len(doc.tables):
                result["warnings"].append(f"Table {t} not found"); continue
            table = doc.tables[t]
            if r >= len(table.rows):
                result["warnings"].append(f"Table {t} row {r} not found"); continue
            if c >= len(table.rows[r].cells):
                result["warnings"].append(f"Table {t} R{r}C{c} not found"); continue
            cell = table.rows[r].cells[c]

            # 跳过标签单元格
            if inspect_data:
                for ti in inspect_data.get("tables", []):
                    for ci in ti.get("cells", []):
                        if ci["row"] == r and ci["column"] == c and ci.get("is_label"):
                            result["warnings"].append(
                                f"跳过标签单元格 R{r}C{c}: \"{ci.get('text','')}\"")
                            break

            fill_utils.fill_cell_safe(
                cell, entry["text"],
                font_name=entry.get("font_name", fill_utils.FONT_BODY),
                font_size_pt=entry.get("font_size", fill_utils.SIZE_BODY),
                bold=entry.get("bold", False),
                east_asia=entry.get("east_asia", entry.get("font_name")),
                align=entry.get("align"),
            )
            result["cells_filled"] += 1

        doc.save(output_path)
        result["success"] = True
    except Exception as e:
        result["error"] = f"直接填充失败: {e}"
    return result


def verify_original_unchanged(template_path: Path, original_hash: str) -> bool:
    with open(template_path, 'rb') as f:
        return hashlib.sha256(f.read()).hexdigest() == original_hash


def _insert_image_at_match(doc: Document, match_text: str, image_path: str | None, caption: str):
    """Find a paragraph containing match_text and insert an image(or placeholder) after it."""
    for para in doc.paragraphs:
        if match_text and match_text in para.text:
            fill_utils.insert_image_or_placeholder(para, image_path, caption)
            break


def fill_by_roles(template_source: Path, roles_data: dict, roles_map_path: Path,
                   output_path: Path) -> dict:
    """按角色名填充模板。

    1. 从 roles_map (extract_template 输出) 获取每个角色的 cell 坐标
    2. 复制 template_source 为 output
    3. 用 roles_data 中对应 key 的值替换 fillable cell
    4. 标签 cell 保持不变
    """
    if not HAS_DOCX:
        return {"error": "Missing dependencies: python-docx, docxtpl"}

    result = {"success": False, "output": str(output_path), "roles_filled": [], "warnings": []}

    try:
        with open(roles_map_path, 'r', encoding='utf-8') as f:
            roles_map = json.load(f)
    except Exception as e:
        result["error"] = f"无法读取角色映射: {e}"
        return result

    try:
        shutil.copy(template_source, output_path)
        doc = Document(output_path)

        filled = 0
        for role_name, targets in roles_map.get("roles", {}).items():
            if role_name not in roles_data:
                continue
            value = roles_data[role_name]
            for target in targets:
                t_idx = target["table"]
                r_idx = target["row"]
                c_idx = target["col"]
                fmt = target.get("format", {})

                if t_idx >= len(doc.tables):
                    continue
                table = doc.tables[t_idx]
                if r_idx >= len(table.rows):
                    continue
                if c_idx >= len(table.rows[r_idx].cells):
                    continue
                cell = table.rows[r_idx].cells[c_idx]

                fill_utils.fill_cell_safe(
                    cell, str(value),
                    font_name=fmt.get("font_name", fill_utils.FONT_BODY),
                    font_size_pt=fmt.get("font_size", fill_utils.SIZE_BODY),
                    bold=fmt.get("bold", False),
                    east_asia=fmt.get("east_asia"),
                    align=fmt.get("alignment"),
                )
                result["roles_filled"].append(role_name)
                filled += 1

        doc.save(output_path)
        result["success"] = True
    except Exception as e:
        result["error"] = f"角色填充失败: {e}"

    return result


# Backward-compatible alias for tests
fill_template = fill_with_inspect


def main():
    parser = argparse.ArgumentParser(description='Fill DOCX template')
    parser.add_argument('--template', '-t', required=True, help='模板文件 (.docx/.doc)')
    parser.add_argument('--output', '-o', required=True, help='输出文件路径')
    parser.add_argument('--data', '-d', help='JSON 数据文件（placeholder 模式）')
    parser.add_argument('--cells', help='JSON 单元格直填文件（无 placeholder 模式）')
    parser.add_argument('--roles', help='按角色名填充的 data JSON（配合 --template-source）')
    parser.add_argument('--template-source', help='成品报告作为模板来源（配合 --roles）')
    parser.add_argument('--inspect', help='inspect_template.py 输出的 JSON')
    parser.add_argument('--images', help='JSON 图片插入配置（按段落文本匹配插入位置）')
    parser.add_argument('--style', choices=['perfect', 'normal'], default='normal')
    args = parser.parse_args()

    template_path = Path(args.template)
    output_path = Path(args.output)

    if not template_path.exists():
        print(f"Error: Template not found: {template_path}", file=sys.stderr); sys.exit(1)

    with open(template_path, 'rb') as f:
        original_hash = hashlib.sha256(f.read()).hexdigest()

    inspect_data = None
    if args.inspect:
        ip = Path(args.inspect)
        if ip.exists():
            with open(ip, 'r', encoding='utf-8') as f:
                inspect_data = json.load(f)

    # 判断模式
    if args.roles:
        # 角色名填充模式（从成品报告提取模板结构）
        roles_path = Path(args.roles)
        if not roles_path.exists():
            print(f"Error: Roles file not found: {roles_path}", file=sys.stderr); sys.exit(1)
        if not args.template_source:
            print("Error: --roles 需要 --template-source（成品报告路径）", file=sys.stderr); sys.exit(1)
        template_source = Path(args.template_source)
        if not template_source.exists():
            print(f"Error 模板来源不存在: {template_source}", file=sys.stderr); sys.exit(1)
        with open(roles_path, 'r', encoding='utf-8') as f:
            roles_data = json.load(f)
        result = fill_by_roles(template_source, roles_data, roles_path, output_path)

    elif args.cells:
        cells_path = Path(args.cells)
        if not cells_path.exists():
            print(f"Error: Cells file not found: {cells_path}", file=sys.stderr); sys.exit(1)
        with open(cells_path, 'r', encoding='utf-8') as f:
            cells_data = json.load(f)
        result = fill_cells_direct(template_path, cells_data, output_path, inspect_data)
    else:
        if not args.data:
            print("Error: 需要 --data（placeholder 模式）或 --cells（直接填充模式）", file=sys.stderr)
            sys.exit(1)
        data_path = Path(args.data)
        if not data_path.exists():
            print(f"Error: Data file not found: {data_path}", file=sys.stderr); sys.exit(1)
        result = fill_with_inspect(template_path, data_path, output_path, inspect_data)

    if not verify_original_unchanged(template_path, original_hash):
        print("Error: Original template was modified!", file=sys.stderr); sys.exit(1)

    # ═══ 图片插入（后处理） ═══
    if args.images and result.get("success"):
        images_path = Path(args.images)
        if images_path.exists():
            with open(images_path, 'r', encoding='utf-8') as f:
                images_config = json.load(f)
            # images_config: [{"match": "文本片段", "path": "screenshots/x.jpg", "caption": "图注"}, ...]
            out_doc = Document(output_path)
            for img_cfg in images_config:
                match_text = img_cfg.get("match", "")
                img_path = img_cfg.get("path")
                caption = img_cfg.get("caption", "此处插入照片")
                _insert_image_at_match(out_doc, match_text, img_path, caption)
            out_doc.save(output_path)

            # Output screenshot manifest for manual verification
            manifest = []
            for img_cfg in images_config:
                manifest.append({
                    "caption": img_cfg.get("caption", ""),
                    "source_path": img_cfg.get("path", ""),
                    "insert_location": img_cfg.get("match", ""),
                    "需要人工验证": True
                })
            manifest_path = output_path.with_name(output_path.stem + '_manifest.json')
            with open(str(manifest_path), 'w', encoding='utf-8') as f:
                json.dump(manifest, f, indent=2, ensure_ascii=False)

    if result.get("warnings"):
        for w in result["warnings"]:
            print(f"WARNING: {w}", file=sys.stderr)

    print(json.dumps(result, indent=2, ensure_ascii=False))
    sys.exit(0 if result.get("success") else 1)


if __name__ == '__main__':
    main()
