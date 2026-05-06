#!/usr/bin/env python3
"""Validate generated DOCX report — structural and formatting integrity check."""

import argparse
import json
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def validate(input_path: Path, inspect_data: dict = None, strict: bool = False, image_config_path: str = None) -> dict:
    """Validate a generated DOCX report. Returns check results."""
    if not HAS_DOCX:
        return {"valid": False, "checks": [], "warnings": ["Missing dependency: python-docx"]}

    result = {
        "valid": True,
        "checks": [],
        "warnings": []
    }

    try:
        doc = Document(str(input_path))
        result["checks"].append({"name": "file_openable", "status": "PASS"})
    except Exception as e:
        result["valid"] = False
        result["checks"].append({"name": "file_openable", "status": "FAIL", "detail": str(e)})
        return result  # Can't continue if file won't open

    # Check 2: No residual {{placeholders}}
    try:
        full_text = " ".join([p.text for p in doc.paragraphs])
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text += " " + cell.text
        remaining = re.findall(r'\{\{[^}]+\}\}', full_text)
        if remaining:
            result["valid"] = False if strict else result["valid"]
            result["checks"].append({
                "name": "no_remaining_placeholders",
                "status": "FAIL" if strict else "WARN",
                "detail": f"Found {len(remaining)} unreplaced placeholders: {remaining}"
            })
        else:
            result["checks"].append({"name": "no_remaining_placeholders", "status": "PASS"})
    except Exception as e:
        result["warnings"].append(f"Placeholder check failed: {e}")

    # Check 3: Table structure matches inspect data
    if inspect_data and "tables" in inspect_data:
        for t_idx, table_spec in enumerate(inspect_data.get("tables", [])):
            if t_idx >= len(doc.tables):
                result["valid"] = False
                result["checks"].append({
                    "name": f"table_{t_idx}_exists",
                    "status": "FAIL",
                    "detail": f"Inspect expects table {t_idx}, but output has only {len(doc.tables)} tables"
                })
                continue

            out_table = doc.tables[t_idx]
            expected_rows = table_spec.get("rows", 0)
            actual_rows = len(out_table.rows)
            if actual_rows != expected_rows:
                result["valid"] = False
                result["checks"].append({
                    "name": f"table_{t_idx}_rows",
                    "status": "FAIL",
                    "detail": f"Expected {expected_rows} rows, got {actual_rows}"
                })
            else:
                result["checks"].append({"name": f"table_{t_idx}_rows", "status": "PASS"})

    # Check 4: Label cells preserved (if inspect data available)
    if inspect_data:
        for t_idx, table_spec in enumerate(inspect_data.get("tables", [])):
            if t_idx >= len(doc.tables):
                continue
            out_table = doc.tables[t_idx]
            for cell_spec in table_spec.get("cells", []):
                if not cell_spec.get("is_label"):
                    continue
                row, col = cell_spec["row"], cell_spec["column"]
                if row >= len(out_table.rows):
                    continue
                if col >= len(out_table.rows[row].cells):
                    continue
                expected_text = cell_spec.get("text", "")
                actual_text = out_table.rows[row].cells[col].text.strip()
                if actual_text == "":
                    # Label might have been overwritten
                    result["warnings"].append(
                        f"Label R{row}C{col} was '{expected_text}' but now empty"
                    )
                elif actual_text != expected_text:
                    result["warnings"].append(
                        f"Label R{row}C{col} overwritten: '{expected_text}' -> '{actual_text}'"
                    )

    if result["warnings"] and strict:
        result["valid"] = False

    # Check 5: Image insertion
    if image_config_path:
        check = _check_image_insertion(doc, image_config_path)
        result["checks"].append(check)
        if check["status"] == "WARN" and strict:
            result["valid"] = False

    return result


def _check_image_insertion(doc, image_config_path: str) -> dict:
    """Verify all configured images were inserted into the document."""
    try:
        with open(image_config_path, 'r', encoding='utf-8') as f:
            configs = json.load(f)
    except Exception:
        return {"name": "image_insertion", "status": "SKIP", "detail": "Cannot read image config"}

    total = len(configs)
    matched = []
    unmatched = []

    for cfg in configs:
        caption = cfg.get("caption", "")
        found = False
        for p in doc.paragraphs:
            if caption and caption in p.text:
                found = True
                break
        if found:
            matched.append(caption)
        else:
            unmatched.append(cfg.get("path", "unknown"))

    if unmatched and total > 0:
        return {
            "name": "image_insertion",
            "status": "WARN",
            "detail": f"{len(matched)}/{total} images matched, unmatched: {unmatched}"
        }
    return {"name": "image_insertion", "status": "PASS", "detail": f"All {total} images matched"}


def main():
    parser = argparse.ArgumentParser(description='Validate generated DOCX report')
    parser.add_argument('--input', '-i', required=True, help='Generated DOCX report')
    parser.add_argument('--inspect', help='Inspect JSON from inspect_template.py')
    parser.add_argument('--strict', action='store_true', help='Fail on warnings too')
    parser.add_argument('--images', help='Image config JSON (image-config.json) for insertion verification')
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Error: File not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    inspect_data = None
    if args.inspect:
        inspect_path = Path(args.inspect)
        if inspect_path.exists():
            with open(inspect_path, 'r', encoding='utf-8') as f:
                inspect_data = json.load(f)

    image_config_path = args.images
    result = validate(input_path, inspect_data, args.strict, image_config_path)

    print(json.dumps(result, indent=2, ensure_ascii=False))
    sys.exit(0 if result["valid"] else 1)


if __name__ == '__main__':
    main()
