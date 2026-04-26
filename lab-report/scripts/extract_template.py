#!/usr/bin/env python3
"""从成品实验报告中提取模板结构 — 角色名→cell坐标映射。

输入：一份已填好的同课程报告（.docx 或 .doc）
输出：JSON 映射文件，记录每个单元格的角色名、格式、是否可填

原理：
  对应行/列的 "label" 单元格（如 "课程名称"）→ 其相邻 "fillable" 单元格 → 角色名就是那个 label。
  合并单元格 → 使用 table.rows[r].cells[c] 避免 grid 索引偏移。
"""

import argparse
import json
import re
import sys
from pathlib import Path

try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# 常见标签文本 → 标准化角色名
ROLE_ALIASES = {
    "课程名称": "课程名称",
    "课程": "课程名称",
    "课程代码": "课程代码",
    "任课教师": "任课教师",
    "教师": "任课教师",
    "指导教师": "任课教师",
    "学生姓名": "学生姓名",
    "姓名": "学生姓名",
    "年级": "专业年级",
    "专业年级": "专业年级",
    "专业": "专业年级",
    "班级": "专业年级",
    "学号": "学号",
    "实验名称": "实验名称",
    "实验项目": "实验名称",
    "实验类型": "实验类型",
    "实验学时": "实验学时",
    "实验日期": "实验日期",
    "实验地点": "实验地点",
    "实验环境": "实验环境",
    "实验设备": "实验设备",
    "提交文档": "提交文档说明",
}


def _normalize_role(text: str) -> str:
    """Standardize label text to a known role name."""
    t = text.strip().replace("：", "").replace(":", "").replace(" ", "").replace("\n", "")
    for key, value in ROLE_ALIASES.items():
        if key in t:
            return value
    return t


def _is_label_cell(text: str) -> bool:
    """判断一个单元格是标签（如'课程名称'）还是内容（如'工程伦理'）"""
    t = text.strip()
    if not t:
        return False
    # 标签通常是短的、纯中文的，末尾不带标点（或只有冒号）
    if len(t) > 20:
        return False
    if t.startswith(("{{", "（")):
        return False
    return True


def extract(template_path: Path) -> dict:
    """从成品报告提取模板结构。"""
    if not HAS_DOCX:
        return {"error": "Missing dependency: python-docx"}

    result = {
        "source": template_path.name,
        "tables": [],
        "roles": {},  # role_name → list of (table, row, col) targets
        "body_sections": [],
    }

    try:
        doc = Document(template_path)

        for t_idx, table in enumerate(doc.tables):
            table_info = {"index": t_idx, "rows": len(table.rows),
                          "cells": []}

            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if not text:
                        continue

                    # 获取格式
                    font_name = None
                    font_size = None
                    bold = None
                    east_asia = None
                    alignment = None

                    first_para = cell.paragraphs[0]
                    if first_para.alignment is not None:
                        alignment = str(first_para.alignment)
                    for run in first_para.runs:
                        if run.font.name:
                            font_name = run.font.name
                        if run.font.size:
                            font_size = round(run.font.size / 12700, 1)
                        if run.font.bold is not None:
                            bold = bool(run.font.bold)
                        rPr = run._element.find(
                            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr')
                        if rPr is not None:
                            rFonts = rPr.find(
                                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
                            if rFonts is not None:
                                ea = rFonts.get(
                                    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia')
                                if ea:
                                    east_asia = ea
                        break

                    # 判断类型
                    is_label = _is_label_cell(text)
                    role = _normalize_role(text) if is_label else None

                    cell_info = {
                        "row": r_idx,
                        "col": c_idx,
                        "text": text,
                        "is_label": is_label,
                        "role": role,
                        "format": {
                            "font_name": font_name or "宋体",
                            "font_size": font_size or 12,
                            "bold": bold or False,
                            "east_asia": east_asia,
                            "alignment": alignment,
                        }
                    }

                    # 尝试从同行相邻 label 推断角色
                    if not is_label:
                        # 找同行最近的 label cell
                        for peer_c_idx, peer_cell in enumerate(row.cells):
                            if peer_c_idx == c_idx:
                                continue
                            peer_text = peer_cell.text.strip()
                            if _is_label_cell(peer_text):
                                cell_info["inferred_role"] = _normalize_role(peer_text)
                                # 注册到 roles 映射
                                inferred = cell_info["inferred_role"]
                                if inferred not in result["roles"]:
                                    result["roles"][inferred] = []
                                result["roles"][inferred].append({
                                    "table": t_idx, "row": r_idx, "col": c_idx,
                                    "format": cell_info["format"],
                                    "text_preview": text,
                                })
                                break

                    table_info["cells"].append(cell_info)
            result["tables"].append(table_info)

        if not result["roles"]:
            result["warning"] = (
                "未检测到可映射的角色。请确认模板是'标签+值'邻列布局。"
                "如果所有单元格都是标签（整表已填满），请尝试使用另一份空白未填的报告。"
            )

    except Exception as e:
        result["error"] = str(e)

    return result


def print_summary(result: dict):
    """人类可读输出。"""
    if "error" in result:
        print(f"Error: {result['error']}")
        return

    print(f"模板来源: {result['source']}")
    print(f"检测到 {len(result['tables'])} 个表格")
    print(f"角色映射: {len(result['roles'])} 个")
    print("-" * 40)

    for role, targets in result["roles"].items():
        print(f"\n  {role}:")
        for t in targets:
            print(f"    Table {t['table']} R{t['row']}C{t['col']}: 现值为 \"{t['text_preview'][:40]}\"")
            print(f"    格式: {t['format']['font_name']} {t['format']['font_size']}pt"
                  f"  {'B' if t['format']['bold'] else 'R'}"
                  f"  eastAsia={t['format']['east_asia']}"
                  f"  align={t['format']['alignment']}")

    if result.get("warning"):
        print(f"\n⚠️  {result['warning']}")


def main():
    parser = argparse.ArgumentParser(
        description='从成品报告提取模板结构（角色→cell映射）'
    )
    parser.add_argument('--input', '-i', required=True, help='成品报告 .docx/.doc')
    parser.add_argument('--format', '-f', choices=['json', 'human'], default='json')
    parser.add_argument('--output', '-o', help='保存到文件')
    args = parser.parse_args()

    filepath = Path(args.input)
    if not filepath.exists():
        print(f"Error: File not found: {filepath}", file=sys.stderr)
        sys.exit(1)

    result = extract(filepath)

    if args.format == 'human':
        print_summary(result)
        output_str = ""
    else:
        output_str = json.dumps(result, indent=2, ensure_ascii=False)
        if not args.output:
            print(output_str)

    if args.output:
        Path(args.output).write_text(output_str or json.dumps(result, indent=2, ensure_ascii=False),
                                     encoding='utf-8')
        print(f"Saved to {args.output}")


if __name__ == '__main__':
    main()
