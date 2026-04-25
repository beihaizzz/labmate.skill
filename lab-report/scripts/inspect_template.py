#!/usr/bin/env python3
"""Inspect a DOCX template and dump complete run-level formatting.

This is the MANDATORY first step before any fill operation.
Run it, read the output, THEN write your fill data — never guess formatting.
"""

import argparse
import json
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

PLACEHOLDER_RE = re.compile(r'\{\{[^}]+\}\}')

def _get_run_font_info(run) -> dict:
    """Extract font information from a single run."""
    info = {
        "text_preview": run.text[:80] if run.text else "",
        "font_name": None,
        "font_size_emu": None,  # EMU units, convert to pt: EMU / 12700
        "font_size_pt": None,
        "bold": None,
        "italic": None,
        "underline": None,
        "color_rgb": None,
        "east_asia": None,  # <w:eastAsia> attribute
    }

    try:
        if run.font.name:
            info["font_name"] = run.font.name
        if run.font.size:
            info["font_size_emu"] = run.font.size
            info["font_size_pt"] = round(run.font.size / 12700, 1)
        if run.font.bold is not None:
            info["bold"] = bool(run.font.bold)
        if run.font.italic is not None:
            info["italic"] = bool(run.font.italic)
        if run.font.underline is not None:
            info["underline"] = run.font.underline
        if run.font.color and run.font.color.rgb:
            info["color_rgb"] = str(run.font.color.rgb)

        # Read eastAsia from XML directly (not available via python-docx API)
        rPr = run._element.find(qn('w:rPr'))
        if rPr is not None:
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                ea = rFonts.get(qn('w:eastAsia'))
                if ea:
                    info["east_asia"] = ea
    except Exception:
        pass

    return info


def _get_paragraph_info(para, para_idx: int) -> dict:
    """Get paragraph-level formatting."""
    info = {
        "index": para_idx,
        "text_preview": para.text[:120] if para.text else "",
        "text_length": len(para.text.strip()),
        "style_name": para.style.name if para.style else None,
        "alignment": str(para.alignment) if para.alignment else None,
        "has_placeholder": bool(PLACEHOLDER_RE.search(para.text)),
        "runs": [],
    }

    for run in para.runs:
        info["runs"].append(_get_run_font_info(run))

    return info


def _get_cell_info(cell, row_idx: int, col_idx: int) -> dict:
    """Get cell-level formatting, distinguishing labels from placeholders."""
    cell_text = cell.text.strip()
    has_placeholder = bool(PLACEHOLDER_RE.search(cell_text))

    info = {
        "row": row_idx,
        "column": col_idx,
        "text": cell_text,
        "text_length": len(cell_text),
        "has_placeholder": has_placeholder,
        "is_label": bool(cell_text and not has_placeholder),
        "paragraphs": [],
    }

    for p_idx, para in enumerate(cell.paragraphs):
        info["paragraphs"].append(_get_paragraph_info(para, p_idx))

    return info


def inspect_template(filepath: Path) -> dict:
    """Inspect a DOCX template and return complete formatting info."""
    if not HAS_DOCX:
        return {"error": "Missing dependency: python-docx"}

    result = {
        "filename": filepath.name,
        "has_placeholders": False,
        "tables": [],
        "body_paragraphs": [],
        "summary": {
            "total_tables": 0,
            "total_paragraphs": 0,
            "placeholder_count": 0,
            "label_cells": 0,
        },
    }

    try:
        doc = Document(filepath)

        # ─── Inspect tables ───
        for t_idx, table in enumerate(doc.tables):
            table_info = {
                "index": t_idx,
                "rows": len(table.rows),
                "columns": len(table.columns) if table.columns else 0,
                "cells": [],
            }
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    ci = _get_cell_info(cell, r_idx, c_idx)
                    table_info["cells"].append(ci)

                    if ci["has_placeholder"]:
                        result["has_placeholders"] = True
                        result["summary"]["placeholder_count"] += 1
                    if ci["is_label"]:
                        result["summary"]["label_cells"] += 1

            result["tables"].append(table_info)
            result["summary"]["total_tables"] += 1

        # ─── Inspect body paragraphs ───
        for p_idx, para in enumerate(doc.paragraphs):
            if not para.text.strip():
                continue
            pi = _get_paragraph_info(para, p_idx)
            result["body_paragraphs"].append(pi)
            result["summary"]["total_paragraphs"] += 1

            if pi["has_placeholder"]:
                result["has_placeholders"] = True
                result["summary"]["placeholder_count"] += 1

    except Exception as e:
        result["error"] = str(e)

    return result


def print_human_readable(result: dict):
    """Print a summary that the AI can read to understand the template."""
    if "error" in result:
        print(f"Error: {result['error']}")
        return

    print("=" * 60)
    print(f"TEMPLATE INSPECTION: {result['filename']}")
    print("=" * 60)
    print(f"Tables: {result['summary']['total_tables']}")
    print(f"Body paragraphs: {result['summary']['total_paragraphs']}")
    print(f"Placeholder count: {result['summary']['placeholder_count']}")
    print(f"Label cells (preserve!): {result['summary']['label_cells']}")
    print()

    for table in result["tables"]:
        print(f"\n─── Table {table['index']} ({table['rows']}×{table['columns']}) ───")

        # Group by row
        current_row = None
        for cell in table["cells"]:
            if cell["row"] != current_row:
                if current_row is not None:
                    print("  ──")
                current_row = cell["row"]

            tag = "📌 LABEL" if cell["is_label"] else ("📝 PLACEHOLDER" if cell["has_placeholder"] else "  empty ")
            print(f"  R{cell['row']}C{cell['column']} [{tag}] \"{cell['text'][:50]}\"")

            for p in cell["paragraphs"]:
                align_str = f" align={p['alignment']}" if p.get("alignment") else ""
                for r in p["runs"]:
                    if r["text_preview"]:
                        pt = r["font_size_pt"] if r["font_size_pt"] else "?"
                        ea = r["east_asia"] if r["east_asia"] else "—"
                        b = "B" if r["bold"] else " "
                        print(f"    font={r['font_name']} {pt}pt eastAsia={ea} [{b}]{align_str} \"{r['text_preview'][:40]}\"")

    if result["body_paragraphs"]:
        print(f"\n─── Body Paragraphs ({len(result['body_paragraphs'])}) ───")
        for p in result["body_paragraphs"]:
            tag = "📝" if p["has_placeholder"] else "  "
            print(f"  {tag} \"{p['text_preview'][:60]}\"")
            for r in p["runs"]:
                pt = r["font_size_pt"] if r["font_size_pt"] else "?"
                ea = r["east_asia"] if r["east_asia"] else "—"
                b = "B" if r["bold"] else " "
                al = f" align={p['alignment']}" if p.get("alignment") else ""
                print(f"    font={r['font_name']} {pt}pt eastAsia={ea} [{b}]{al} \"{r['text_preview'][:40]}\"")

    print("\n" + "=" * 60)
    print("HOW TO USE THIS DATA")
    print("=" * 60)
    print("1. Each [LABEL] cell must be PRESERVED as-is — do NOT overwrite.")
    print("2. Each [PLACEHOLDER] cell can be replaced — use the SAME font/size/eastAsia/align as shown.")
    print("3. eastAsia column: if '—' (none), do NOT add eastAsia to that cell.")
    print("4. ALIGN column: match this exactly (CENTER/LEFT/RIGHT). Table 0 value cells should be CENTER.")
    print("5. Font sizes in pt — use exactly these values.")
    print("6. Build data JSON with keys matching each {{placeholder}}.")
    print("7. Body paragraphs (成段叙述) should have first-line indent of 24pt (2 chars).")
    print("8. List items (1., 2., (1)) should NOT have indent.")


def main():
    parser = argparse.ArgumentParser(
        description='Inspect DOCX template formatting (MANDATORY before fill)'
    )
    parser.add_argument('--input', '-i', required=True, help='Template .docx file')
    parser.add_argument('--format', '-f', choices=['json', 'human'], default='json',
                        help='json=full data (feed to fill_template); human=readable summary')
    parser.add_argument('--output', '-o', help='Save output to file instead of stdout')
    args = parser.parse_args()

    filepath = Path(args.input)
    if not filepath.exists():
        print(f"Error: File not found: {filepath}", file=sys.stderr)
        sys.exit(1)

    result = inspect_template(filepath)

    output_str = ""
    if args.format == 'human':
        # human format prints directly to stdout
        import io
        buf = io.StringIO()
        old_out = sys.stdout
        sys.stdout = buf
        try:
            print_human_readable(result)
        finally:
            sys.stdout = old_out
        output_str = buf.getvalue()
        print(output_str, end='')
    else:
        output_str = json.dumps(result, indent=2, ensure_ascii=False)
        if not args.output:
            print(output_str)

    if args.output:
        Path(args.output).write_text(output_str, encoding='utf-8')
        print(f"Saved to {args.output}")

    sys.exit(0 if "error" not in result else 1)


if __name__ == '__main__':
    main()
