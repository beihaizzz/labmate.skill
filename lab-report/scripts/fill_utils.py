"""Formatting utilities for lab-report DOCX template filling.

Shared constants and helper functions so all scripts use consistent formatting.
"""

from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx import Document
from typing import Optional
import re

# ── 格式化常量 ────────────────────────────────────────────────────────────────────
FONT_TITLE = '黑体'
FONT_BODY = '宋体'
FONT_CODE = 'Courier New'

SIZE_TABLE0_VALUE = 14
SIZE_H1 = 12
SIZE_H2 = 11
SIZE_BODY = 12
SIZE_CODE = 9
SIZE_IMAGE_HINT = 10

INDENT_2CHAR = Pt(24)

_LIST_ITEM_PATTERNS = re.compile(
    r'^\s*('
    r'\d+[．\.、]'
    r'|[a-zA-Z][\)\.]'
    r'|[（\(][一二三四五六七八九十\d]+[）\)]'
    r'|[一二三四五六七八九十]+[、．\.]'
    r')\s*'
)


def is_list_item(text: str) -> bool:
    """Check if text looks like a numbered/bulleted list item."""
    return bool(_LIST_ITEM_PATTERNS.match(text.strip()))


def is_body_paragraph(text: str, min_len: int = 20) -> bool:
    """Check if text is a descriptive body paragraph."""
    t = text.strip()
    if not t:
        return False
    if is_list_item(t):
        return False
    if len(t) < min_len:
        return False
    if t.startswith('[') and t.endswith(']'):
        return False
    if t.startswith('{{') and t.endswith('}}'):
        return False
    return True


def set_run_font(run, font_name=None, font_size_pt=None, bold=None, east_asia=None):
    """Set font properties on a run."""
    if font_name:
        run.font.name = font_name
    if font_size_pt:
        run.font.size = Pt(font_size_pt)
    if bold is not None:
        run.font.bold = bold
    if east_asia:
        rPr = run._element.find(qn('w:rPr'))
        if rPr is None:
            rPr = run._element.makeelement(qn('w:rPr'), {})
            run._element.insert(0, rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = rPr.makeelement(qn('w:rFonts'), {})
            rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), east_asia)


def set_paragraph_alignment(para, alignment_str: str | None):
    """Map alignment string from inspect output to WD_ALIGN_PARAGRAPH."""
    if not alignment_str:
        return
    mapping = {
        'CENTER': WD_ALIGN_PARAGRAPH.CENTER,
        'LEFT': WD_ALIGN_PARAGRAPH.LEFT,
        'RIGHT': WD_ALIGN_PARAGRAPH.RIGHT,
        'JUSTIFY': WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    align = mapping.get(alignment_str.upper())
    if align is not None:
        para.alignment = align


def apply_first_line_indent(para, indent=INDENT_2CHAR):
    """Apply first-line indent if paragraph is a body paragraph."""
    text = para.text.strip()
    if is_body_paragraph(text):
        para.paragraph_format.first_line_indent = indent


def add_run(para, text, font_name=FONT_BODY, font_size_pt=SIZE_BODY,
            bold=False, east_asia=None):
    """Add a formatted run to a paragraph."""
    run = para.add_run(text)
    set_run_font(run, font_name=font_name, font_size_pt=font_size_pt,
                 bold=bold, east_asia=east_asia)
    return run


def heading_run(para, text, font_name=FONT_TITLE, font_size_pt=SIZE_H1):
    """Add a heading run."""
    return add_run(para, text, font_name=font_name, font_size_pt=font_size_pt,
                   bold=True, east_asia=font_name)


def body_run(para, text, font_name=FONT_BODY, font_size_pt=SIZE_BODY):
    """Add a body text run."""
    return add_run(para, text, font_name=font_name, font_size_pt=font_size_pt,
                   bold=False, east_asia=font_name)


def image_hint_run(para, label: str):
    """Add a styled image placeholder."""
    run = add_run(para, f"\n[{label}]\n", font_name=FONT_BODY,
                  font_size_pt=SIZE_IMAGE_HINT, bold=False, east_asia=FONT_BODY)
    run.italic = True
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return run


# ── 中文正文段落快捷写入 ──────────────────────────────────────────────────────────

def add_chinese_body_para(cell, text: str, font_name=FONT_BODY, font_size_pt=SIZE_BODY):
    """Write a Chinese body paragraph with first-line indent.
    
    自动清除单元格内容后写入一段中文正文，含首行缩进两声符、宋体、常规。
    """
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    set_run_font(run, font_name=font_name, font_size_pt=font_size_pt,
                 bold=False, east_asia=font_name)
    p.paragraph_format.first_line_indent = INDENT_2CHAR
    return p


# ── 安全单元格查找与写入 ──────────────────────────────────────────────────────────

def find_cell_by_content(table, row_index: int, text_contains: str) -> Optional[object]:
    """Find a cell in a specific row by matching partial text content.
    
    比 table.cell(r, c) 安全，避免合并单元格导致的索引偏移。
    
    Example:
        cell = find_cell_by_content(table, 0, "提交文档")
        if cell:
            fill_cell_safe(...)
    """
    if row_index >= len(table.rows):
        return None
    row = table.rows[row_index]
    for cell in row.cells:
        if text_contains in cell.text.strip():
            return cell
    return None


def fill_cell_safe(cell, text: str, font_name=FONT_BODY, font_size_pt=SIZE_BODY,
                   bold=False, east_asia=None, align=None):
    """Safely write text into a cell.
    
    清除旧内容、写入新文本、设字体/对齐/eastAsia。支持:
    - 普通填充: fill_cell_safe(cell, "张三")
    - 表格值填充: fill_cell_safe(cell, "张三", font_name=FONT_BODY, align="CENTER")
    """
    for p in cell.paragraphs:
        p.clear()
    first_para = cell.paragraphs[0]
    run = first_para.add_run(text)
    set_run_font(run, font_name=font_name, font_size_pt=font_size_pt,
                 bold=bold, east_asia=east_asia)
    if align:
        set_paragraph_alignment(first_para, align)
    return first_para


# ── 合并单元格检测 ─────────────────────────────────────────────────────────────────

def get_cell_grid_range(table, row_idx: int, col_idx: int) -> dict:
    """Detect if a cell is merged and return its grid span.
    
    Returns: {colspan, rowspan} or None if not merged.
    """
    row = table.rows[row_idx]
    if col_idx >= len(row.cells):
        return {"colspan": 1, "rowspan": 1}
    tc = row.cells[col_idx]._tc
    tc_pr = tc.find(qn('w:tcPr'))
    if tc_pr is None:
        return {"colspan": 1, "rowspan": 1}
    grid_span = tc_pr.find(qn('w:gridSpan'))
    v_merge = tc_pr.find(qn('w:vMerge'))
    colspan = int(grid_span.get(qn('w:val'))) if grid_span is not None else 1
    rowspan = 2 if v_merge is not None else 1
    return {"colspan": colspan, "rowspan": rowspan}


def iter_cells_merged_aware(table):
    """Iterate cells with merged-cell handling.
    
    Yields (row_idx, col_idx, cell, grid_info) for each logical cell position.
    Skips cells that are hidden by a vertical merge.
    """
    grid = get_table_grid(table)
    for r in range(len(table.rows)):
        for c in range(len(grid[r])):
            if grid[r][c] is None:
                continue  # hidden by merge
            cell = grid[r][c]
            info = get_cell_grid_range(table, r, c)
            yield r, c, cell, info


def get_table_grid(table):
    """Build a (rows × max_cols) grid of cell references.
    
    Cells hidden by vertical merge are set to None.
    """
    rows = []
    for row in table.rows:
        row_cells = list(row.cells)
        rows.append(row_cells)
    return rows
