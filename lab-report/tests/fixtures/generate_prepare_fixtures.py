"""Generate 4 synthetic blank template DOCX files for testing auto_prepare_template.py."""
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
import json
import os

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
FIXTURES_DIR = SCRIPT_DIR


def set_font(run, name, size_pt, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size_pt)
    run.font.bold = bold


def create_standard(doc, output_path):
    """4×4 info grid table, labels in odd cols (0,2), empty value cells in even cols (1,3)."""
    doc.add_paragraph()  # simple spacer
    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"

    labels = [
        ("课程名称", "任课教师"),
        ("学生姓名", "学号"),
        ("专业年级", "实验日期"),
        ("实验名称", "实验类型"),
    ]

    for row_idx, (label_left, label_right) in enumerate(labels):
        # Left label
        cell = table.rows[row_idx].cells[0]
        p = cell.paragraphs[0]
        run = p.add_run(label_left)
        set_font(run, "SimSun", 12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Left value (empty)
        cell = table.rows[row_idx].cells[1]
        p = cell.paragraphs[0]
        run = p.add_run("")
        set_font(run, "SimSun", 14)

        # Right label
        cell = table.rows[row_idx].cells[2]
        p = cell.paragraphs[0]
        run = p.add_run(label_right)
        set_font(run, "SimSun", 12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Right value (empty)
        cell = table.rows[row_idx].cells[3]
        p = cell.paragraphs[0]
        run = p.add_run("")
        set_font(run, "SimSun", 14)

    doc.save(output_path)
    print(f"Created: {output_path}")


def create_hints(doc, output_path):
    """Same 4×4 layout but even cols contain hint text."""
    doc.add_paragraph()
    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"

    labels = [
        ("课程名称", "任课教师"),
        ("学生姓名", "学号"),
        ("专业年级", "实验日期"),
        ("实验名称", "实验类型"),
    ]
    # Hints for value cells in reading order (col 1 then col 3 per row)
    hints = [
        ("_____", "（请填写）"),
        ("待定", "_____"),
    ]

    for row_idx, (label_left, label_right) in enumerate(labels):
        # Left label
        cell = table.rows[row_idx].cells[0]
        p = cell.paragraphs[0]
        run = p.add_run(label_left)
        set_font(run, "SimSun", 12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Left value with hint
        hint_text = hints[row_idx % 2][0]
        cell = table.rows[row_idx].cells[1]
        p = cell.paragraphs[0]
        run = p.add_run(hint_text)
        set_font(run, "SimSun", 14)

        # Right label
        cell = table.rows[row_idx].cells[2]
        p = cell.paragraphs[0]
        run = p.add_run(label_right)
        set_font(run, "SimSun", 12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Right value with hint
        hint_text = hints[row_idx % 2][1]
        cell = table.rows[row_idx].cells[3]
        p = cell.paragraphs[0]
        run = p.add_run(hint_text)
        set_font(run, "SimSun", 14)

    doc.save(output_path)
    print(f"Created: {output_path}")


def create_merged(doc, output_path):
    """Table with merged header row spanning 4 cols, label-value pairs below."""
    doc.add_paragraph()
    table = doc.add_table(rows=5, cols=4)
    table.style = "Table Grid"

    # Row 0: Merged header "实验报告封面"
    cell = table.rows[0].cells[0]
    # Merge cells 0-3 in row 0
    cell.merge(table.rows[0].cells[3])
    p = cell.paragraphs[0]
    run = p.add_run("实验报告封面")
    set_font(run, "SimHei", 14, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Row 1: 课程名称 | (empty) | 任课教师 | (empty)
    row_data = [
        ("课程名称", ""),
        ("任课教师", ""),
    ]
    for col_idx, (label, value) in enumerate(row_data):
        lcell = table.rows[1].cells[col_idx * 2]
        p = lcell.paragraphs[0]
        run = p.add_run(label)
        set_font(run, "SimSun", 12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        vcell = table.rows[1].cells[col_idx * 2 + 1]
        p = vcell.paragraphs[0]
        run = p.add_run(value)
        set_font(run, "SimSun", 14)

    # Row 2: 学生姓名 | (empty) | 学号 | (empty)
    row_data = [
        ("学生姓名", ""),
        ("学号", ""),
    ]
    for col_idx, (label, value) in enumerate(row_data):
        lcell = table.rows[2].cells[col_idx * 2]
        p = lcell.paragraphs[0]
        run = p.add_run(label)
        set_font(run, "SimSun", 12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        vcell = table.rows[2].cells[col_idx * 2 + 1]
        p = vcell.paragraphs[0]
        run = p.add_run(value)
        set_font(run, "SimSun", 14)

    # Row 3: 专业年级 | (empty) | 实验日期 | (empty)
    row_data = [
        ("专业年级", ""),
        ("实验日期", ""),
    ]
    for col_idx, (label, value) in enumerate(row_data):
        lcell = table.rows[3].cells[col_idx * 2]
        p = lcell.paragraphs[0]
        run = p.add_run(label)
        set_font(run, "SimSun", 12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        vcell = table.rows[3].cells[col_idx * 2 + 1]
        p = vcell.paragraphs[0]
        run = p.add_run(value)
        set_font(run, "SimSun", 14)

    # Row 4: 实验名称 | (empty) | 实验类型 | (empty)
    row_data = [
        ("实验名称", ""),
        ("实验类型", ""),
    ]
    for col_idx, (label, value) in enumerate(row_data):
        lcell = table.rows[4].cells[col_idx * 2]
        p = lcell.paragraphs[0]
        run = p.add_run(label)
        set_font(run, "SimSun", 12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        vcell = table.rows[4].cells[col_idx * 2 + 1]
        p = vcell.paragraphs[0]
        run = p.add_run(value)
        set_font(run, "SimSun", 14)

    doc.save(output_path)
    print(f"Created: {output_path}")


def create_nonstandard(doc, output_path):
    """Labels NOT in ROLE_ALIASES: 课题组, 提交日期, 指导教师 with empty value cells."""
    doc.add_paragraph()
    table = doc.add_table(rows=3, cols=4)
    table.style = "Table Grid"

    labels = [
        ("课题组", "提交日期"),
        ("指导教师", "审阅教师"),
        ("报告日期", "成绩"),
    ]

    for row_idx, (label_left, label_right) in enumerate(labels):
        # Left label
        cell = table.rows[row_idx].cells[0]
        p = cell.paragraphs[0]
        run = p.add_run(label_left)
        set_font(run, "SimSun", 12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Left value (empty)
        cell = table.rows[row_idx].cells[1]
        p = cell.paragraphs[0]
        run = p.add_run("")
        set_font(run, "SimSun", 14)

        # Right label
        cell = table.rows[row_idx].cells[2]
        p = cell.paragraphs[0]
        run = p.add_run(label_right)
        set_font(run, "SimSun", 12)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Right value (empty)
        cell = table.rows[row_idx].cells[3]
        p = cell.paragraphs[0]
        run = p.add_run("")
        set_font(run, "SimSun", 14)

    doc.save(output_path)
    print(f"Created: {output_path}")


def main():
    os.makedirs(FIXTURES_DIR, exist_ok=True)

    create_standard(Document(), os.path.join(FIXTURES_DIR, "prepare_standard.docx"))
    create_hints(Document(), os.path.join(FIXTURES_DIR, "prepare_hints.docx"))
    create_merged(Document(), os.path.join(FIXTURES_DIR, "prepare_merged.docx"))
    create_nonstandard(Document(), os.path.join(FIXTURES_DIR, "prepare_nonstandard.docx"))

    print("\nAll 4 fixtures created successfully.")


if __name__ == "__main__":
    main()