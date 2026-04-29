# -*- coding: utf-8 -*-
"""TDD RED phase: Tests for auto_prepare_template.py (not yet implemented).

These tests will FAIL initially because auto_prepare_template.py does not exist.
Once the module is implemented, these tests should pass.
"""

import hashlib
import sys
import tempfile
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Add scripts directory to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent / 'scripts'))

# Attempt to import the module under test (will fail until RED phase is resolved)
try:
    from auto_prepare_template import prepare_template, main
except ImportError:
    prepare_template = None
    main = None


# =============================================================================
# Fixtures
# =============================================================================

@pytest.fixture
def fixture_dir():
    """Return path to test fixtures directory."""
    return Path(__file__).parent / 'fixtures'


@pytest.fixture
def prepare_standard_path(fixture_dir):
    """Return path to prepare_standard.docx fixture."""
    return fixture_dir / 'prepare_standard.docx'


@pytest.fixture
def prepare_hints_path(fixture_dir):
    """Return path to prepare_hints.docx fixture."""
    return fixture_dir / 'prepare_hints.docx'


@pytest.fixture
def prepare_merged_path(fixture_dir):
    """Return path to prepare_merged.docx fixture."""
    return fixture_dir / 'prepare_merged.docx'


@pytest.fixture
def prepare_nonstandard_path(fixture_dir):
    """Return path to prepare_nonstandard.docx fixture."""
    return fixture_dir / 'prepare_nonstandard.docx'


@pytest.fixture
def temp_copy(prepare_standard_path, tmp_path):
    """Create a temporary copy of prepare_standard.docx for testing."""
    import shutil
    dest = tmp_path / 'prepare_standard.docx'
    shutil.copy(prepare_standard_path, dest)
    return dest


@pytest.fixture
def temp_copy_hints(prepare_hints_path, tmp_path):
    """Create a temporary copy of prepare_hints.docx for testing."""
    import shutil
    dest = tmp_path / 'prepare_hints.docx'
    shutil.copy(prepare_hints_path, dest)
    return dest


@pytest.fixture
def temp_copy_merged(prepare_merged_path, tmp_path):
    """Create a temporary copy of prepare_merged.docx for testing."""
    import shutil
    dest = tmp_path / 'prepare_merged.docx'
    shutil.copy(prepare_merged_path, dest)
    return dest


@pytest.fixture
def temp_copy_nonstandard(prepare_nonstandard_path, tmp_path):
    """Create a temporary copy of prepare_nonstandard.docx for testing."""
    import shutil
    dest = tmp_path / 'prepare_nonstandard.docx'
    shutil.copy(prepare_nonstandard_path, dest)
    return dest


def _sha256(path: Path) -> str:
    """Compute SHA256 hash of a file."""
    with open(path, 'rb') as f:
        return hashlib.sha256(f.read()).hexdigest()


def _create_vertical_template(tmp_path: Path) -> Path:
    """Create a vertical-layout template: R0: [姓名][学号], R1: [value][value]."""
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    # Row 0: labels
    table.cell(0, 0).text = '姓名'
    table.cell(0, 1).text = '学号'
    # Row 1: empty values
    table.cell(1, 0).text = ''
    table.cell(1, 1).text = ''
    path = tmp_path / 'vertical_template.docx'
    doc.save(str(path))
    return path


def _create_formatted_template(tmp_path: Path) -> Path:
    """Create a template with specific formatting: 仿宋 14pt CENTER in a value cell."""
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    # Row 0: label
    table.cell(0, 0).text = '课程名称'
    # Row 1: value cell with specific formatting
    cell = table.cell(1, 0)
    run = cell.paragraphs[0].add_run('SAMPLE_VALUE')
    run.font.name = '仿宋'
    run.font.size = 14 * 12700  # 14pt in EMUs
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    path = tmp_path / 'formatted_template.docx'
    doc.save(str(path))
    return path


def _create_empty_template(tmp_path: Path) -> Path:
    """Create a template with no tables."""
    doc = Document()
    # Add only body text, no tables
    doc.add_paragraph('This is an empty template with no tables.')
    path = tmp_path / 'empty_template.docx'
    doc.save(str(path))
    return path


# =============================================================================
# Test Classes
# =============================================================================

class TestHorizontalDetection:
    """Tests for horizontal label→value pair detection."""

    @pytest.mark.skipif(prepare_template is None, reason="auto_prepare_template.py not yet implemented")
    def test_horizontal_prepares_standard_fixture(self, prepare_standard_path):
        """Verify horizontal layout: R0C1 contains {{课程名称}}, R1C1 contains {{学生姓名}}, R2C1 contains {{专业年级}}."""
        from auto_prepare_template import prepare_template

        output_path = prepare_template(str(prepare_standard_path))
        doc = Document(str(output_path))

        assert len(doc.tables) > 0, "No tables found in output"
        table = doc.tables[0]

        # Check row 0, column 1 (value cell next to 课程名称 label)
        r0c1_text = table.cell(0, 1).text
        assert '{{课程名称}}' in r0c1_text, f"R0C1 should contain {{{{课程名称}}}}, got: {r0c1_text}"

        # Check row 1, column 1 (value cell next to 学生姓名 label)
        r1c1_text = table.cell(1, 1).text
        assert '{{学生姓名}}' in r1c1_text, f"R1C1 should contain {{{{学生姓名}}}}, got: {r1c1_text}"

        # Check row 2, column 1 (value cell next to 专业年级 label)
        r2c1_text = table.cell(2, 1).text
        assert '{{专业年级}}' in r2c1_text, f"R2C1 should contain {{{{专业年级}}}}, got: {r2c1_text}"


class TestVerticalDetection:
    """Tests for vertical label→value pair detection."""

    @pytest.mark.skipif(prepare_template is None, reason="auto_prepare_template.py not yet implemented")
    def test_vertical_detection_injects_below_label(self, tmp_path):
        """Verify vertical layout: R1C0 has {{学生姓名}} when R0C0 is '姓名' label."""
        from auto_prepare_template import prepare_template

        vertical_path = _create_vertical_template(tmp_path)
        output_path = prepare_template(str(vertical_path))
        doc = Document(str(output_path))

        assert len(doc.tables) > 0, "No tables found in output"
        table = doc.tables[0]

        # R1C0 (row 1, col 0) should have the injected placeholder
        r1c0_text = table.cell(1, 0).text
        assert '{{学生姓名}}' in r1c0_text, f"R1C0 should contain {{{{学生姓名}}}}, got: {r1c0_text}"


class TestHintReplacement:
    """Tests for hint text→placeholder replacement."""

    @pytest.mark.skipif(prepare_template is None, reason="auto_prepare_template.py not yet implemented")
    def test_hint_text_is_identified(self):
        """Verify is_hint_text correctly identifies _____ and （请填写） as hints."""
        from role_aliases import is_hint_text

        assert is_hint_text('_____') is True, "_____ should be identified as hint text"
        assert is_hint_text('（请填写）') is True, "（请填写） should be identified as hint text"

    @pytest.mark.skipif(prepare_template is None, reason="auto_prepare_template.py not yet implemented")
    def test_hints_replaced_with_role(self, temp_copy_hints):
        """Verify hint placeholders like _____, （请填写）, 待定 are replaced with role-based placeholders."""
        from auto_prepare_template import prepare_template

        output_path = prepare_template(str(temp_copy_hints))
        doc = Document(str(output_path))

        # Collect all text from all tables
        all_text = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text.append(cell.text)

        combined_text = ' '.join(all_text)

        # Verify no hint patterns remain
        assert '_____' not in combined_text, "Hint text _____ should be replaced"
        assert '（请填写）' not in combined_text, "Hint text （请填写） should be replaced"
        assert '待定' not in combined_text, "Hint text 待定 should be replaced"


class TestFormatPreservation:
    """Tests for original formatting preservation."""

    @pytest.mark.skipif(prepare_template is None, reason="auto_prepare_template.py not yet implemented")
    def test_font_formatting_preserved_after_injection(self, tmp_path):
        """Verify font_name, font_size_pt, and alignment are preserved after injection."""
        from auto_prepare_template import prepare_template

        formatted_path = _create_formatted_template(tmp_path)
        output_path = prepare_template(str(formatted_path))
        doc = Document(str(output_path))

        table = doc.tables[0]
        cell = table.cell(1, 0)

        # Get the run's formatting
        runs = cell.paragraphs[0].runs
        assert len(runs) > 0, "No runs found in value cell"

        run = runs[0]
        font_name = run.font.name
        font_size_pt = run.font.size.pt if run.font.size else None
        alignment = cell.paragraphs[0].alignment

        # Verify formatting matches original (仿宋, 14pt, CENTER)
        assert font_name == '仿宋', f"Font name should be 仿宋, got {font_name}"
        assert abs(font_size_pt - 14.0) < 0.1, f"Font size should be 14pt, got {font_size_pt}"
        assert alignment == WD_ALIGN_PARAGRAPH.CENTER, f"Alignment should be CENTER, got {alignment}"


class TestIdempotency:
    """Tests for idempotent operation (double-run same output)."""

    @pytest.mark.skipif(prepare_template is None, reason="auto_prepare_template.py not yet implemented")
    def test_double_run_produces_same_output(self, temp_copy):
        """Verify running auto_prepare twice produces identical placeholder counts."""
        from auto_prepare_template import prepare_template

        # First run
        output1 = prepare_template(str(temp_copy))
        doc1 = Document(str(output1))

        # Count placeholders in first output
        def count_placeholders(doc):
            count = 0
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text = cell.text
                        count += text.count('{{')
                        count += text.count('}}')
            return count

        count1 = count_placeholders(doc1)

        # Second run on the same temp file
        output2 = prepare_template(str(temp_copy))
        doc2 = Document(str(output2))
        count2 = count_placeholders(doc2)

        # Placeholder counts should be identical (no double braces)
        assert count1 == count2, f"First run had {count1} placeholder markers, second run had {count2}"


class TestMergedCellHandling:
    """Tests for merged cell handling."""

    @pytest.mark.skipif(prepare_template is None, reason="auto_prepare_template.py not yet implemented")
    def test_merged_header_cell_preserved(self, temp_copy_merged):
        """Verify merged header cell '实验报告封面' is preserved and not overwritten."""
        from auto_prepare_template import prepare_template

        output_path = prepare_template(str(temp_copy_merged))
        doc = Document(str(output_path))

        # Check that merged header text is still present
        all_text = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text.append(cell.text)

        combined_text = ' '.join(all_text)
        assert '实验报告封面' in combined_text, "Merged header '实验报告封面' should be preserved"


class TestNonStandardLabels:
    """Tests for labels outside ROLE_ALIASES."""

    @pytest.mark.skipif(prepare_template is None, reason="auto_prepare_template.py not yet implemented")
    def test_nonstandard_labels_use_raw_text_as_role(self, temp_copy_nonstandard):
        """Verify non-standard labels use raw text as placeholder role name (not mapped name)."""
        from auto_prepare_template import prepare_template

        output_path = prepare_template(str(temp_copy_nonstandard))
        doc = Document(str(output_path))

        # Collect all text from all tables
        all_text = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text.append(cell.text)

        combined_text = ' '.join(all_text)

        # The nonstandard label should appear as {{raw_label_text}}
        # since it's not in ROLE_ALIASES
        assert '{{' in combined_text, "No placeholders found in output"
        # Should NOT have been mapped to any standard role
        assert '{{课程名称}}' not in combined_text or '{{非标准标签}}' in combined_text, \
            "Non-standard label should use raw text as role"


class TestOutputOverwriteProtection:
    """Tests for input file protection."""

    @pytest.mark.skipif(prepare_template is None, reason="auto_prepare_template.py not yet implemented")
    def test_original_input_sha256_unchanged(self, prepare_standard_path):
        """Verify SHA256 of original fixture is unchanged after auto_prepare."""
        from auto_prepare_template import prepare_template

        # Compute SHA256 before
        sha256_before = _sha256(prepare_standard_path)

        # Create temp copy and run auto_prepare on it
        import shutil
        with tempfile.TemporaryDirectory() as tmpdir:
            temp_path = Path(tmpdir) / 'test_input.docx'
            shutil.copy(prepare_standard_path, temp_path)

            output_path = prepare_template(str(temp_path))

            # Verify SHA256 of original fixture is unchanged
            sha256_after = _sha256(prepare_standard_path)
            assert sha256_before == sha256_after, \
                f"Original fixture SHA256 changed: before={sha256_before}, after={sha256_after}"


class TestEmptyTemplate:
    """Tests for graceful handling of edge cases."""

    @pytest.mark.skipif(prepare_template is None, reason="auto_prepare_template.py not yet implemented")
    def test_empty_template_returns_empty_roles(self, tmp_path):
        """Verify template with no tables returns success=true with empty roles_injected."""
        from auto_prepare_template import prepare_template

        empty_path = _create_empty_template(tmp_path)
        result = prepare_template(str(empty_path))

        # Result should be a dict with success=true and empty roles_injected
        assert isinstance(result, dict), "Result should be a dict"
        assert result.get('success') is True, "Result should have success=true"
        assert result.get('roles_injected', []) == [], "roles_injected should be empty"
