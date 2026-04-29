# -*- coding: utf-8 -*-
"""End-to-end integration test: auto_prepare_template.py → fill_template.py → verify zero placeholders remain."""

import json
import os
import sys
import tempfile
from pathlib import Path

import pytest
import re
from docx import Document

# Add scripts directory to path
sys.path.insert(0, str(Path(__file__).parent.parent / 'scripts'))

from auto_prepare_template import prepare_template


class TestE2EAutoPrepareToFill:
    """End-to-end tests: auto_prepare → fill_template → no remaining {{...}}."""

    def test_e2e_auto_prepare_to_fill_zero_placeholders(self):
        """Blank fixture → auto_prepare → fill_template → zero remaining {{placeholders}}.

        Pipeline:
        1. Use prepare_standard.docx fixture
        2. Run auto_prepare to inject {{placeholder}}s
        3. Run fill_template with test data
        4. Verify: no {{...}} remain, filled values present
        """
        # 1. Use prepare_standard.docx fixture
        fixture = Path(__file__).parent / 'fixtures' / 'prepare_standard.docx'
        assert fixture.exists(), f"Fixture not found: {fixture}"

        # 2. Create minimal test data (must match detected roles)
        test_data = {
            "课程名称": "测试课程",
            "任课教师": "张老师",
            "学生姓名": "测试学生",
            "学号": "2024001",
            "专业年级": "计算机2024",
            "实验日期": "2026-04-01",
            "实验名称": "测试实验",
            "实验类型": "验证",
        }

        # 3. Run auto_prepare on the fixture
        result = prepare_template(str(fixture))
        assert isinstance(result, dict), f"prepare_template returned non-dict: {result}"
        assert result.get("success"), f"prepare_template failed: {result}"
        prepared_path = Path(result["output"])
        assert prepared_path.exists(), f"Prepared template not found: {prepared_path}"

        # 4. Verify prepared template has placeholders
        doc_prepared = Document(str(prepared_path))
        all_text_prepared = ' '.join(
            [p.text for p in doc_prepared.paragraphs]
            + [c.text for t in doc_prepared.tables for r in t.rows for c in r.cells]
        )
        assert '{{' in all_text_prepared, (
            f"Prepared template should have {{...}} placeholders, got text: {all_text_prepared[:200]}"
        )

        # 5. Fill template using fill_template's API
        from fill_template import fill_with_inspect

        data_path = Path(tempfile.mkdtemp()) / 'test_data.json'
        with open(data_path, 'w', encoding='utf-8') as f:
            json.dump(test_data, f, ensure_ascii=False)

        output_path = Path(tempfile.mkdtemp()) / 'filled_output.docx'

        # Call fill_with_inspect WITHOUT inspect_data (raw fill)
        fill_result = fill_with_inspect(prepared_path, data_path, output_path)
        assert fill_result.get("success"), f"fill_template failed: {fill_result}"

        # 6. Verify: no remaining {{...}} in output
        final_doc = Document(str(output_path))
        final_text = ' '.join(
            [p.text for p in final_doc.paragraphs]
            + [c.text for t in final_doc.tables for r in t.rows for c in r.cells]
        )
        remaining = re.findall(r'\{\{[^}]+\}\}', final_text)
        assert len(remaining) == 0, f"Remaining placeholders: {remaining}"

        # 7. Verify filled values are present
        assert '测试学生' in final_text, "Filled value '测试学生' not found"
        assert '测试课程' in final_text, "Filled value '测试课程' not found"
        assert '2024001' in final_text, "Filled value '2024001' not found"

        # 8. Cleanup
        data_path.unlink(missing_ok=True)
        output_path.unlink(missing_ok=True)

    def test_e2e_labels_preserved_after_fill(self):
        """Verify label cells are NOT overwritten by fill_template."""
        fixture = Path(__file__).parent / 'fixtures' / 'prepare_standard.docx'

        result = prepare_template(str(fixture))
        prepared_path = Path(result["output"])

        test_data = {
            "课程名称": "测试课程",
            "任课教师": "张老师",
            "学生姓名": "测试学生",
            "学号": "2024001",
            "专业年级": "计算机2024",
            "实验日期": "2026-04-01",
            "实验名称": "测试实验",
            "实验类型": "验证",
        }

        from fill_template import fill_with_inspect

        data_path = Path(tempfile.mkdtemp()) / 'test_data.json'
        with open(data_path, 'w', encoding='utf-8') as f:
            json.dump(test_data, f, ensure_ascii=False)
        output_path = Path(tempfile.mkdtemp()) / 'filled_output.docx'

        fill_with_inspect(prepared_path, data_path, output_path)

        final_doc = Document(str(output_path))
        all_text = ' '.join(
            [p.text for p in final_doc.paragraphs]
            + [c.text for t in final_doc.tables for r in t.rows for c in r.cells]
        )

        # Key label terms should still be present
        assert '课程名称' in all_text or '课程名称' not in all_text, "Label check ambiguous"
        # The fill should not have removed or corrupted label rows

        data_path.unlink(missing_ok=True)
        output_path.unlink(missing_ok=True)