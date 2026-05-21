"""Tests for validate_docx.py — TDD RED phase for long_line_check."""
import pytest
from docx import Document

from scripts.validate_docx import validate


class TestLongLineCheck:
    """Tests for detecting over-long continuous ASCII lines in DOCX content."""

    def test_long_line_detected(self, tmp_path):
        """Long continuous ASCII (250 chars) triggers WARN status."""
        doc = Document()
        # 250 continuous ASCII with no spaces
        long_text = "A" * 250
        doc.add_paragraph(long_text)
        docx_path = tmp_path / "test_long_line.docx"
        doc.save(str(docx_path))

        result = validate(str(docx_path))

        checks = result.get("checks", [])
        long_line_checks = [c for c in checks if c.get("name") == "long_line_check"]
        assert len(long_line_checks) > 0, "long_line_check should be present"
        assert long_line_checks[0].get("status") == "WARN", \
            "long_line_check status should be WARN for 250-char ASCII line"

    def test_normal_report_no_warning(self, tmp_path):
        """Normal text (Chinese + short ASCII with spaces) should have PASS check."""
        doc = Document()
        doc.add_paragraph("这是一个正常的实验报告段落，包含中文内容。")
        doc.add_paragraph("设置GPIOA_Pin_0为输入模式，初始化完成。")
        docx_path = tmp_path / "test_normal.docx"
        doc.save(str(docx_path))

        result = validate(str(docx_path))

        checks = result.get("checks", [])
        long_line_checks = [c for c in checks if c.get("name") == "long_line_check"]
        # long_line_check must exist and be PASS for normal text
        assert len(long_line_checks) > 0, "long_line_check should be present"
        assert long_line_checks[0].get("status") == "PASS", \
            "normal text should have long_line_check PASS"

    def test_mixed_chinese_no_false_positive(self, tmp_path):
        """Short ASCII snippets in Chinese context should have PASS (not WARN)."""
        doc = Document()
        doc.add_paragraph("实验步骤：设置GPIOA_Pin_0为输入模式，配置PB1为输出模式。")
        doc.add_paragraph("数据采集完成后，分析结果并记录实验数据。")
        docx_path = tmp_path / "test_mixed.docx"
        doc.save(str(docx_path))

        result = validate(str(docx_path))

        checks = result.get("checks", [])
        long_line_checks = [c for c in checks if c.get("name") == "long_line_check"]
        # long_line_check must exist and be PASS (not WARN)
        assert len(long_line_checks) > 0, "long_line_check should be present"
        assert long_line_checks[0].get("status") != "WARN", \
            "mixed Chinese should have PASS not WARN"

    def test_long_ascii_in_table_cell(self, tmp_path):
        """Long ASCII in table cell also triggers WARN."""
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        cell.text = "B" * 200
        docx_path = tmp_path / "test_table_long.docx"
        doc.save(str(docx_path))

        result = validate(str(docx_path))

        checks = result.get("checks", [])
        long_line_checks = [c for c in checks if c.get("name") == "long_line_check"]
        assert len(long_line_checks) > 0, "long_line_check should be present for table cell"
        assert long_line_checks[0].get("status") == "WARN", \
            "long ASCII in table cell should trigger WARN"

    def test_threshold_exact_200(self, tmp_path):
        """Boundary: 199 chars PASS, 200 chars WARN."""
        # 199 chars — should PASS (not WARN)
        doc_pass = Document()
        doc_pass.add_paragraph("C" * 199)
        docx_path_pass = tmp_path / "test_199.docx"
        doc_pass.save(str(docx_path_pass))

        result_pass = validate(str(docx_path_pass))
        checks_pass = result_pass.get("checks", [])
        llc_pass = [c for c in checks_pass if c.get("name") == "long_line_check"]
        assert len(llc_pass) > 0, "long_line_check should be present for 199-char line"
        assert llc_pass[0].get("status") == "PASS", \
            "199-char line should have PASS status"

        # 200 chars — should WARN
        doc_warn = Document()
        doc_warn.add_paragraph("D" * 200)
        docx_path_warn = tmp_path / "test_200.docx"
        doc_warn.save(str(docx_path_warn))

        result_warn = validate(str(docx_path_warn))
        checks_warn = result_warn.get("checks", [])
        llc_warn = [c for c in checks_warn if c.get("name") == "long_line_check"]
        assert len(llc_warn) > 0, "long_line_check should be present for 200-char line"
        assert llc_warn[0].get("status") == "WARN", \
            "200-char line should trigger WARN"