#!/usr/bin/env python3
"""Generate default lab report DOCX template."""

from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

def set_cjk_font(run, font_name, size=None):
    """Set CJK font for a run using w:eastAsia attribute."""
    run.font.name = font_name
    run.font.size = size
    # Set East Asia font
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rPr.insert(0, rFonts)

def add_heading(doc, text, level=1):
    """Add a heading with CJK font."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    if level == 1:
        set_cjk_font(run, '黑体', Pt(16))
        run.bold = True
    else:
        set_cjk_font(run, '黑体', Pt(14))
        run.bold = True
    return p

def add_body(doc, text):
    """Add body text with CJK font."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_cjk_font(run, '宋体', Pt(12))
    return p

def add_placeholder_paragraph(doc, placeholder):
    """Add a placeholder line for user input."""
    p = doc.add_paragraph()
    run = p.add_run(placeholder)
    set_cjk_font(run, '宋体', Pt(12))
    run.italic = True
    p.paragraph_format.space_after = Pt(6)
    return p

def create_template():
    doc = Document()

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('实验报告')
    set_cjk_font(run, '黑体', Pt(22))
    run.bold = True

    doc.add_paragraph()  # spacing

    # Info table
    info_table = doc.add_table(rows=5, cols=4)
    info_table.style = 'Table Grid'

    fields = [
        ['姓名', '{{姓名}}', '学号', '{{学号}}'],
        ['学院', '{{学院}}', '专业', '{{专业}}'],
        ['班级', '{{班级}}', '课程名', '{{课程名}}'],
        ['实验名称', '{{实验名称}}', '', ''],
        ['实验日期', '{{实验日期}}', '实验地点', '{{实验地点}}'],
    ]

    for i, row_data in enumerate(fields):
        row = info_table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = cell_text
            for para in cell.paragraphs:
                for run in para.runs:
                    set_cjk_font(run, '宋体', Pt(11))

    doc.add_paragraph()  # spacing

    # Section: 实验目的
    add_heading(doc, '一、实验目的')
    add_placeholder_paragraph(doc, '{{实验目的}}')

    # Section: 实验原理
    add_heading(doc, '二、实验原理')
    add_placeholder_paragraph(doc, '{{实验原理}}')

    # Section: 实验器材
    add_heading(doc, '三、实验器材')
    add_placeholder_paragraph(doc, '{{实验器材}}')

    # Section: 实验步骤
    add_heading(doc, '四、实验步骤')
    add_placeholder_paragraph(doc, '{{实验步骤}}')

    # Section: 实验数据
    add_heading(doc, '五、实验数据')
    add_placeholder_paragraph(doc, '{{实验数据}}')

    # Section: 实验结果
    add_heading(doc, '六、实验结果')
    add_placeholder_paragraph(doc, '{{实验结果}}')

    # Section: 实验结论
    add_heading(doc, '七、实验结论')
    add_placeholder_paragraph(doc, '{{实验结论}}')

    doc.save('lab-report/assets/report_template.docx')
    print('Created: lab-report/assets/report_template.docx')

if __name__ == '__main__':
    create_template()